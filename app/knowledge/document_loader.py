from __future__ import annotations

import csv
import json
from dataclasses import dataclass
from pathlib import Path


SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".csv", ".json", ".docx"}


@dataclass
class LoadedDocument:
    path: Path
    text: str
    metadata: dict


class DocumentLoader:
    def __init__(self, folder: Path):
        self.folder = folder

    def load(self) -> list[LoadedDocument]:
        self.folder.mkdir(parents=True, exist_ok=True)
        docs: list[LoadedDocument] = []
        for path in sorted(self.folder.rglob("*")):
            if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS:
                text = self._read(path)
                if text.strip():
                    docs.append(LoadedDocument(path=path, text=text, metadata={"file_name": path.name}))
        return docs

    def _read(self, path: Path) -> str:
        suffix = path.suffix.lower()
        if suffix in {".txt", ".md"}:
            return path.read_text(encoding="utf-8", errors="ignore")
        if suffix == ".json":
            data = json.loads(path.read_text(encoding="utf-8", errors="ignore"))
            return json.dumps(data, indent=2, ensure_ascii=False)
        if suffix == ".csv":
            rows: list[str] = []
            with path.open("r", encoding="utf-8", errors="ignore", newline="") as handle:
                for row in csv.reader(handle):
                    rows.append(", ".join(row))
            return "\n".join(rows)
        if suffix == ".pdf":
            try:
                from pypdf import PdfReader
            except Exception as exc:
                return f"PDF skipped because pypdf is unavailable: {exc}"
            reader = PdfReader(str(path))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        if suffix == ".docx":
            try:
                import docx
            except Exception as exc:
                return f"DOCX skipped because python-docx is unavailable: {exc}"
            document = docx.Document(str(path))
            return "\n".join(paragraph.text for paragraph in document.paragraphs)
        return ""
